from __future__ import annotations

import json
import os
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "作品集标题与描述编辑清单.docx"
COPY_PATH = ROOT / "src" / "data" / "work-card-copy.json"
CONTENT_PATHS = (
    ROOT / "src" / "data" / "portfolio-content.json",
    ROOT / "src" / "data" / "film-works.json",
    ROOT / "src" / "data" / "photo-albums.json",
)
DELETED_SLUG = "film-ranking-analysis"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle, object_pairs_hook=OrderedDict)


def canonical_titles() -> dict[str, str]:
    titles: dict[str, str] = {}
    for path in CONTENT_PATHS:
        payload = load_json(path)
        for work in payload["works"]:
            titles[work["slug"]] = work["title"]
    return titles


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return

    preferred = next((run for run in runs if run.text.strip()), runs[0])
    for run in runs:
        run.text = ""
    preferred.text = text


def set_field_cell(cell, text: str, target_index: int) -> None:
    paragraphs = cell.paragraphs
    if target_index >= len(paragraphs):
        target_index = 0
    for index, paragraph in enumerate(paragraphs):
        replace_paragraph_text(
            paragraph,
            ("\u00a0" + text) if index == target_index and text else "\u00a0",
        )


def paragraph_text(element, document) -> str:
    return Paragraph(element, document).text


def table_from_element(element, document) -> Table:
    return Table(element, document)


def locate_work_blocks(document):
    body = document._element.body
    children = list(body.iterchildren())
    blocks = {}

    for index, element in enumerate(children):
        if element.tag != qn("w:p"):
            continue
        match = re.fullmatch(r"作品标识：(.+)", paragraph_text(element, document).strip())
        if not match:
            continue

        slug = match.group(1)
        table_index = next(
            candidate
            for candidate in range(index - 1, -1, -1)
            if children[candidate].tag == qn("w:tbl")
        )
        heading_index = table_index - 1
        category_index = heading_index - 1
        page_break_index = category_index - 1

        blocks[slug] = {
            "page_break": children[page_break_index],
            "category": children[category_index],
            "heading": children[heading_index],
            "table": children[table_index],
            "identifier": element,
        }
    return blocks


def main() -> None:
    copy = load_json(COPY_PATH)
    titles = canonical_titles()
    document = Document(DOCX_PATH)
    blocks = locate_work_blocks(document)

    expected_slugs = list(copy.keys())
    actual_slugs = [slug for slug in blocks if slug != DELETED_SLUG]
    if actual_slugs != expected_slugs:
        raise RuntimeError(
            "Document/work-card order mismatch:\n"
            f"document={actual_slugs}\nwebsite={expected_slugs}"
        )
    if DELETED_SLUG not in blocks:
        raise RuntimeError(f"Deleted work block not found: {DELETED_SLUG}")

    # Remove the deleted review's entire one-page block while retaining the
    # following page break that introduces the first photography page.
    body = document._element.body
    for element in blocks[DELETED_SLUG].values():
        body.remove(element)

    # The cover summary is the only free-standing paragraph that mentions the
    # previous work count.
    for paragraph in document.paragraphs:
        if "共 30 个作品" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                paragraph.text.replace("共 30 个作品", "共 29 个作品"),
            )

    total = len(expected_slugs)
    for number, slug in enumerate(expected_slugs, start=1):
        block = blocks[slug]
        title = titles[slug]
        fields = copy[slug]

        category_paragraph = Paragraph(block["category"], document)
        category = re.sub(r"\s*·\s*\d{2}/\d{2}\s*$", "", category_paragraph.text)
        replace_paragraph_text(
            category_paragraph,
            f"{category}  ·  {number:02d}/{total:02d}",
        )

        replace_paragraph_text(Paragraph(block["heading"], document), title)

        table = table_from_element(block["table"], document)
        if len(table.rows) < 4 or len(table.columns) < 2:
            raise RuntimeError(f"Unexpected worksheet table geometry for {slug}")

        set_field_cell(table.rows[0].cells[1], title, 0)
        set_field_cell(table.rows[1].cells[1], fields.get("overview", ""), 1)
        set_field_cell(table.rows[2].cells[1], fields.get("description", ""), 3)

    temp_path = DOCX_PATH.with_name(DOCX_PATH.stem + ".syncing.docx")
    document.save(temp_path)
    os.replace(temp_path, DOCX_PATH)

    # Reopen and perform content-level acceptance checks after the atomic save.
    verified = Document(DOCX_PATH)
    verified_blocks = locate_work_blocks(verified)
    if list(verified_blocks) != expected_slugs:
        raise RuntimeError("Saved document does not contain the expected 29 works")

    all_text = "\n".join(
        [paragraph.text for paragraph in verified.paragraphs]
        + [cell.text for table in verified.tables for row in table.rows for cell in row.cells]
    )
    if DELETED_SLUG in all_text or "电影榜单重点单品分析" in all_text:
        raise RuntimeError("Deleted review is still present")
    for slug in expected_slugs:
        if titles[slug] not in all_text:
            raise RuntimeError(f"Missing title after sync: {slug}")
        if copy[slug]["description"] not in all_text:
            raise RuntimeError(f"Missing description after sync: {slug}")

    print(f"Synced {len(expected_slugs)} works to {DOCX_PATH}")
    print(f"Tables: {len(verified.tables)} (1 cover + {len(expected_slugs)} works)")


if __name__ == "__main__":
    main()
