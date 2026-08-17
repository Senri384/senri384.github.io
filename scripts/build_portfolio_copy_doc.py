from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "作品集标题与描述编辑清单.docx"

CATEGORY_ORDER = [
    "experience",
    "game-design",
    "game-research",
    "articles",
    "short-films",
    "scripts",
    "reviews",
    "photography",
]

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 101, 114)
INK = RGBColor(25, 28, 34)
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
PALE_YELLOW = "FFF5CC"
BORDER = "D3DAE3"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size: float, *, bold=False, color=INK, name="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    run2 = paragraph.add_run(" 页")
    set_font(run2, 9, color=MUTED)


def read_data() -> tuple[dict[str, str], list[dict]]:
    sources = [
        ROOT / "src" / "data" / "portfolio-content.json",
        ROOT / "src" / "data" / "film-works.json",
        ROOT / "src" / "data" / "photo-albums.json",
    ]
    categories: dict[str, str] = {}
    works: list[dict] = []
    for source in sources:
        data = json.loads(source.read_text(encoding="utf-8"))
        for category in data.get("categories", []):
            categories[category["slug"]] = category["title"]
        works.extend(data.get("works", []))
    ordered = sorted(
        works,
        key=lambda work: (
            CATEGORY_ORDER.index(work["category"]),
            next(i for i, item in enumerate(works) if item["slug"] == work["slug"]),
        ),
    )
    return categories, ordered


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1)
    run = p.add_run("作品集卡带页文案编辑清单")
    set_font(run, 9, bold=True, color=MUTED)

    footer = section.footer
    set_paragraph_spacing(footer.paragraphs[0], after=0, line=1)
    add_page_number(footer.paragraphs[0])


def add_cover(doc: Document, work_count: int) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)

    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, after=4, line=1)
    run = kicker.add_run("PORTFOLIO COPY WORKSHEET")
    set_font(run, 10, bold=True, color=BLUE, name="Arial")

    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=8, line=1)
    run = title.add_run("作品集卡带页\n文案编辑清单")
    set_font(run, 28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=24, line=1.25)
    run = subtitle.add_run("标题 / 一行概况 / 长描述 / 现有黄字")
    set_font(run, 13, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_fixed_geometry(table, [1701, 7659])
    entries = [
        ("用途", "集中修改卡带盒浏览页即将使用的作品文案。"),
        ("填写方式", "可直接修改【标题】，并在【概况】与【长描述】的空白区域填写内容。"),
        ("迁移提示", "【当前黄字】仅作原内容参考；这些黄字已从卡带盒页面移除。"),
    ]
    for row, (label, value) in zip(table.rows, entries):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(row.cells[0], PALE_BLUE)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.25)
            run = p.add_run(text)
            set_font(run, 10.5, bold=(idx == 0), color=NAVY if idx == 0 else INK)

    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=22, after=0, line=1.25)
    run = note.add_run(f"共 {work_count} 个作品，按网站当前分类顺序排列；每个作品独立一页，便于逐项修改。")
    set_font(run, 10, color=MUTED)


def add_field_row(table, label: str, value: str, *, fill: str, blank_lines: int = 0) -> None:
    row = table.add_row()
    label_cell, value_cell = row.cells
    label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_fill(label_cell, PALE_BLUE)
    set_cell_fill(value_cell, fill)

    p = label_cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.25)
    run = p.add_run(label)
    set_font(run, 10.5, bold=True, color=NAVY)

    p = value_cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.25)
    run = p.add_run(value)
    set_font(run, 10.5, color=INK)
    for _ in range(blank_lines):
        extra = value_cell.add_paragraph()
        set_paragraph_spacing(extra, after=6, line=1.25)
        extra.add_run("\u00a0")


def add_work_page(doc: Document, category_title: str, work: dict, index: int, total: int) -> None:
    category = doc.add_paragraph()
    set_paragraph_spacing(category, after=4, line=1)
    run = category.add_run(f"{category_title}  ·  {index:02d}/{total:02d}")
    set_font(run, 10, bold=True, color=BLUE)

    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(14)
    run = heading.add_run(work["title"])
    set_font(run, 20, bold=True, color=NAVY)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_fixed_geometry(table, [1701, 7659])

    yellow = " / ".join(work.get("tags", [])[:4])
    add_field_row(table, "【标题】", work["title"], fill="FFFFFF")
    add_field_row(table, "【概况】", "", fill="FFFFFF", blank_lines=2)
    add_field_row(table, "【长描述】", "", fill="FFFFFF", blank_lines=8)
    add_field_row(table, "【当前黄字】", yellow, fill=PALE_YELLOW)

    source = doc.add_paragraph()
    set_paragraph_spacing(source, before=10, after=0, line=1.25)
    run = source.add_run(f"作品标识：{work['slug']}")
    set_font(run, 8.5, color=MUTED, name="Consolas")


def build() -> Path:
    categories, works = read_data()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_cover(doc, len(works))

    for index, work in enumerate(works, start=1):
        doc.add_page_break()
        add_work_page(doc, categories[work["category"]], work, index, len(works))

    doc.core_properties.title = "作品集卡带页文案编辑清单"
    doc.core_properties.subject = "作品标题、概况、长描述及现有黄字编辑清单"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
